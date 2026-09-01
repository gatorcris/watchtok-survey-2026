from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "WatchTok_Survey_Technical_Handoff_V8.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(91, 101, 116)
LIGHT = "E8EEF5"


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_font(p.add_run(text))


def add_steps(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_font(p.add_run(text))


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(header.add_run("WATCHTOK ENTHUSIAST SURVEY  |  TECHNICAL HANDOFF"), size=8.5, bold=True, color=MUTED)
add_page_number(section.footer.paragraphs[0])

spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(78)
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(16)
set_font(kicker.add_run("DEPLOYMENT REFERENCE"), size=10, bold=True, color=BLUE)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
set_font(title.add_run("The 2026 WatchTok\nEnthusiast Survey"), size=29, bold=True, color=DARK_BLUE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(52)
set_font(subtitle.add_run("Technical Handoff V8"), size=16, color=MUTED)

meta = doc.add_table(rows=4, cols=2)
set_table_geometry(meta, [2700, 6660])
for row, (label, value) in zip(meta.rows, [
    ("Build date", "September 1, 2026"),
    ("Deployment", "GitHub Pages"),
    ("Data service", "Supabase — WatchTok 2026 Survey"),
    ("Survey version", "V8"),
]):
    shade_cell(row.cells[0], LIGHT)
    p1 = row.cells[0].paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    set_font(p1.add_run(label), bold=True, color=DARK_BLUE)
    p2 = row.cells[1].paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.add_run(value))

doc.add_page_break()

add_heading(doc, "1. V8 outcome")
p = doc.add_paragraph()
set_font(p.add_run("V8 is a deployment-ready static survey client with the complete frozen V7 instrument and the live Supabase data contract. "), bold=True)
set_font(p.add_run("It is designed for GitHub Pages and TikTok’s in-app mobile browser."))
add_bullets(doc, [
    "All 43 participant-facing questions, with stable research IDs preserved behind the interface.",
    "One question per screen, routed progress, Back navigation, and device-local autosave.",
    "Anonymous Supabase authentication; respondents never create or manage an account.",
    "Partial-response autosave, completion timestamps, referral attribution, and test-mode separation.",
    "Optional follow-up email stored only after completion and separately from research answers.",
])

add_heading(doc, "2. Live configuration")
add_bullets(doc, [
    "Project URL: https://djupcnxgpueafhqcsmud.supabase.co",
    "Research table: survey_responses",
    "Optional follow-up table: contact_optins",
    "Private referral roster: referral_sources",
    "Client survey version: V8",
])
p = doc.add_paragraph()
set_font(p.add_run("Security note. "), bold=True, color=DARK_BLUE)
set_font(p.add_run("The repository contains only the publishable browser key. No service-role or secret key is included. Row-level security and table grants remain the enforcement boundary."))

add_heading(doc, "3. Database contract")
add_heading(doc, "survey_responses", 2)
add_bullets(doc, [
    "Identity: id, owner_id, survey_version, and is_test.",
    "State: status, last_question_id, last_display_question, started_at, updated_at, and completed_at.",
    "Research payload: answers JSONB and referral_code.",
    "Uniqueness: one row per owner_id, survey_version, and is_test combination.",
])
add_heading(doc, "contact_optins", 2)
add_bullets(doc, [
    "Fields: id, email, receive_report, future_research, and consented_at.",
    "No survey-response identifier or owner_id is written to the contact record.",
    "Respondents can insert consent but cannot read, list, update, or delete contact records.",
])

add_heading(doc, "4. Required permission migration")
p = doc.add_paragraph()
set_font(p.add_run("Apply supabase/002_authenticated_client_grants.sql once before deployment. "), bold=True)
set_font(p.add_run("The live smoke test confirmed that anonymous authentication works and RLS policies exist, but the original SQL setup did not grant the authenticated API role table privileges."))
add_bullets(doc, [
    "Grants SELECT, INSERT, and UPDATE on survey_responses to authenticated.",
    "Grants INSERT only on contact_optins to authenticated.",
    "Does not grant DELETE, contact-data reads, or respondent access to referral_sources.",
])

add_heading(doc, "5. Routing and validation")
add_bullets(doc, [
    "Q1 option 1 skips Q2 and derives annual watch spending as $0.",
    "Q37 option 7 skips Q38 through Q43.",
    "Q12 option 1 skips Q22 through Q25 and Q14 through Q16.",
    "Q26 option 14 skips Q27 and Q28.",
    "Q14 through Q16 appear only when Q13 includes option 6.",
    "Display questions 12, 19, 20, and 21 enforce a maximum of three selections.",
    "Conflicting None, not-sure, and prefer-not-to-answer choices are exclusive.",
    "Routed questions are serialized as SKIPPED rather than ordinary missing data.",
])

add_heading(doc, "6. Deployment sequence")
add_steps(doc, [
    "Run the Supabase permission migration in the SQL Editor.",
    "Add the final lowercase creator codes to ALLOWED_REFERRAL_CODES in src/config.js. An empty list temporarily accepts any normalized code.",
    "Run npm run check.",
    "Run npm run test:integration. It creates one clearly marked is_test=true response and does not insert contact data.",
    "Upload the project contents to gatorcris/watchtok-survey-2026 and enable GitHub Pages.",
    "Validate direct, referral, returning-partial, completed, and ?test=1 links on iPhone and Android, including TikTok’s in-app browser.",
])

add_heading(doc, "7. Test coverage")
add_bullets(doc, [
    "Frozen 43-question count and display order.",
    "Referral normalization and optional allow-list enforcement.",
    "Every conditional route and SKIPPED serialization.",
    "Exclusive multi-select choices and three-option limits.",
    "Progress calculation on a heavily routed path.",
    "Live anonymous sign-in, partial insert, owner-only read, and completed-status update.",
])

add_heading(doc, "8. Privacy and operating controls")
add_bullets(doc, [
    "No name or TikTok handle is requested in the research instrument.",
    "Optional email is never placed inside answers JSONB.",
    "The public client cannot read another owner’s response.",
    "Test and production responses are separated by is_test and by browser-storage scope.",
    "Deletion requests route to watchtoksurvey@gmail.com.",
    "Final privacy/retention language and a recoverable fielding backup process remain launch gates.",
])

add_heading(doc, "9. Final launch checklist")
add_bullets(doc, [
    "Database permission migration applied and live integration test passed.",
    "Creator referral allow-list frozen and partner links tested.",
    "Privacy and retention language approved.",
    "iPhone, Android, keyboard, focus, and screen-reader checks completed.",
    "Partial and completed exports reconciled with database totals.",
    "Test records confirmed as is_test=true before creator launch.",
])

doc.core_properties.title = "WatchTok Survey Technical Handoff V8"
doc.core_properties.subject = "Deployment and operations reference"
doc.core_properties.author = "Cris Bjelajac"
doc.core_properties.keywords = "WatchTok, survey, V8, Supabase, GitHub Pages"
doc.save(OUT)
print(OUT)
